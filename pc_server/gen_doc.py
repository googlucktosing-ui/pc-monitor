# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()
style = doc.styles["Normal"]
font = style.font
font.name = "Microsoft YaHei"
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

def h(text, level=1):
    hd = doc.add_heading(text, level=level)
    for r in hd.runs:
        r.font.name = "Microsoft YaHei"
        r.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    return hd

def p(text, bold=False, sz=None, color=None, align=None):
    par = doc.add_paragraph()
    run = par.add_run(text)
    run.font.name = "Microsoft YaHei"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if bold: run.bold = True
    if sz: run.font.size = Pt(sz)
    if color: run.font.color.rgb = RGBColor(*color)
    if align: par.alignment = align
    return par

def tbl(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, hdr in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = hdr
        for par in c.paragraphs:
            for r in par.runs:
                r.bold = True; r.font.size = Pt(9)
                r.font.name = "Microsoft YaHei"
                r.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]
            c.text = str(val)
            for par in c.paragraphs:
                for r in par.runs:
                    r.font.size = Pt(9)
                    r.font.name = "Microsoft YaHei"
                    r.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    return t

# Title
doc.add_paragraph(); doc.add_paragraph()
p("PC Monitor 副屏", bold=True, sz=28, color=(26,26,46), align=WD_ALIGN_PARAGRAPH.CENTER)
p("产品落地分析报告", bold=True, sz=22, color=(37,99,235), align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
p("版本 V1.0  |  2026-05-31", sz=11, color=(107,114,128), align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

# 1. Screen
h("一、屏幕选型分析", 1)
p("现有 1.28\" 圆屏（280×240）作为 PC 副屏尺寸偏小。ESP32 能支持的屏幕受限于 RAM（PSRAM）和 SPI 带宽，以下为可行方案对比：")

h("1.1 推荐方案对比", 2)
tbl(["方案","尺寸","分辨率","驱动 IC","接口","批量价","推荐理由"],[
["① 2.8\" IPS 方屏","2.8\"","320×240","ILI9341","SPI","¥15-20","代码改动最小，ST7789 兼容"],
["② 3.5\" IPS 方屏","3.5\"","480×320","ILI9488","SPI/16bit","¥25-35","信息量翻倍，需 PSRAM"],
["③ 4.0\" 圆屏","4.0\"","480×480","GC9A01A","SPI","¥30-40","颜值高，需 PSRAM"],
["④ 3.95\" 方屏","3.95\"","480×480","ST7701S","SPI","¥35-45","显示面积最大"],
])
p("")
par = doc.add_paragraph()
r = par.add_run("首选推荐：2.8\" IPS 方屏（ILI9341）")
r.bold = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor(37,99,235)
r.font.name = "Microsoft YaHei"; r.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
p("• ST7789 驱动与 ILI9341 指令兼容度 90%，移植工作量约 1-2 小时")
p("• 320x240 分辨率，现有 LVGL 布局几乎不改（调坐标即可）")
p("• ESP32-WROOM-32（无 PSRAM）即可流畅运行")
p("• 桌面视觉比例舒适，成本最低")

h("1.2 LVGL 适配工作量", 2)
tbl(["模块","改动内容","工时"],[
["lv_port.c","修改分辨率，ILI9341 初始化序列","约 1h"],
["app_ui.c","调整布局坐标适配方屏","约 1-2h"],
["主控选择","2.8\" 方案无需改动","-"],
])
doc.add_page_break()

# 2. BOM
h("二、关键器件 BOM 与成本预估", 1)
h("2.1 BOM 清单（批量 100pcs，2.8\" IPS 方案）", 2)
tbl(["#","器件","规格","单价","备注"],[
["1","ESP32-WROOM-32 模组","4MB Flash","¥18","乐鑫/安信可"],
["2","2.8\" TFT IPS 屏","320×240, ILI9341","¥18","含 FPC 排线"],
["3","PCB","双层板 60×40mm","¥4","嘉立创 5pcs 摊薄"],
["4","接插件","Type-C 母座+排针","¥3",""],
["5","电源","AMS1117-3.3+电容","¥2","5V→3.3V"],
["6","电阻电容","0603 若干","¥1",""],
["7","外壳","公模桌面支架壳","¥12","1688 采购"],
["8","包装","彩盒+说明书+Type-C线","¥8",""],
])
p("")
p("BOM 合计（不含外壳包装）：¥46    含外壳包装全套：¥66", bold=True, sz=11)

h("2.2 备选方案：ESP32-S3（带 PSRAM）", 2)
tbl(["对比项","ESP32-WROOM","ESP32-S3-WROOM-1"],[
["模组单价","¥18","¥25"],
["可跑分辨率","320x240 流畅","480x320 流畅"],
["适合屏幕","2.8\" 最稳定","3.5\"~4.0\""],
["BOM 总成本","¥71","¥80"],
])

h("2.3 利润分析", 2)
tbl(["项目","金额"],[
["单台成本（2.8\" 方案）","¥66"],
["建议零售价","¥129 ~ ¥159"],
["单台毛利","¥63 ~ ¥93"],
["毛利率","48% ~ 58%"],
])
doc.add_page_break()

# 3. Product
h("三、产品定义与使用场景", 1)
h("3.1 成品描述", 2)
p("一个 2.8\" 桌面信息副屏，无线 WiFi 连接，Type-C 供电。通电后自动连上 PC，实时显示：")
for item in ["CPU 使用率（数字 + 进度条）","GPU 使用率 + 温度","内存 使用量 / 总量","磁盘 使用量 / 总量","网速 上传 / 下载","系统信息（主机名、OS 版本）","时间 + 连接状态指示"]:
    doc.add_paragraph(item, style="List Bullet")

h("3.2 客户使用流程", 2)
tbl(["步骤","操作","耗时"],[
["1","拆箱：屏幕 + Type-C 线 + 说明书","1 分钟"],
["2","插 Type-C 线到电脑 USB 口","5 秒"],
["3","PC 端下载运行 pc_monitor.exe","2 分钟"],
["4","自动连接，开始显示数据","自动（10 秒）"],
])
p("合计：从拆箱到使用约 3 分钟", bold=True, sz=11)

h("3.3 竞品对比", 2)
tbl(["产品类型","价格","连接方式","优势","劣势"],[
["本产品","¥129-159","WiFi 无线","无线、可自定义、开源","品牌知名度低"],
["AIDA64 USB 副屏","¥60-99","USB","便宜","有线占桌位"],
["AIDA64 HDMI 副屏","¥150-300","HDMI","高刷新率","占 HDMI 口"],
["淘宝成品信息屏","¥80-150","USB/WiFi","到手即用","不能自定义"],
["树莓派副屏","¥200+","网络","功能强","贵、体积大"],
])

h("3.4 核心竞争力", 2)
for title, body in [
("① 无线是最大卖点","市面 ¥60-99 的副屏全是有线的。本产品只需一根电源线，放在桌面任何位置。"),
("② 软件完全自主可控","PC 端 + ESP32 固件全部自研。可随时加功能，用户需求可快速响应。"),
("③ 成本优势明显","BOM ¥46-66，售价 ¥129-159，毛利率 50%+。比 HDMI 副屏便宜一半。"),
("④ 差异化定位","本产品：无线·可自定义·开源·可升级   |   竞品：有线·固定功能·闭源·不可升级"),
]:
    par = doc.add_paragraph()
    r = par.add_run(title)
    r.bold = True; r.font.size = Pt(11)
    r.font.name = "Microsoft YaHei"; r.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    p(body)
doc.add_page_break()

# 4. Plan
h("四、量产落地计划", 1)
tbl(["阶段","动作","时间","投入"],[
["P0 屏幕切换","采购 2.8\" ILI9341 屏 → 移植驱动 → 布局适配","1 周","¥100"],
["P1 体验优化","配网 Portal + PC 端 exe 完善","1 周","纯软件"],
["P2 小批量","画 PCB → 定外壳 → 生产 20 台 → 闲鱼上架","2 周","¥1,600"],
["P3 迭代","收集反馈 → 优化体验 → 批量备货","持续","滚动投入"],
])
p("")
p("首批启动资金：约 ¥1,600（20 台物料 + 外壳 + 包装）", bold=True, sz=11)
p("按售价 ¥129/台：首批 20 台全部售出约 ¥2,580 营收", sz=10)

p("")
par = doc.add_paragraph()
r = par.add_run("PC Monitor 副屏项目 · 产品落地分析报告 V1.0 · 2026-05-31")
r.font.size = Pt(9); r.font.color.rgb = RGBColor(107,114,128)
r.font.name = "Microsoft YaHei"; r.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
par.alignment = WD_ALIGN_PARAGRAPH.CENTER

path = r"E:\CODEX-Pj\ESP32\pc_monitor\产品落地分析报告.docx"
doc.save(path)
print(f"Done: {path}")
print(f"Size: {os.path.getsize(path)/1024:.1f} KB")