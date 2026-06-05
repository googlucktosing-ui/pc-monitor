from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import datetime

doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(10)

# Cover
for _ in range(3):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('PC Monitor \u526f\u5c4f\u9879\u76ee')
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0, 102, 204)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\u786c\u4ef6\u539f\u7406\u56fe\u8bbe\u8ba1\u53c2\u8003\u6307\u5357')
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0, 102, 204)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
today = datetime.date.today().strftime("%Y-%m-%d")
run = p.add_run(f'\u7248\u672c: Rev 1.0\n\u65e5\u671f: {today}\n\u4e3b\u63a7: ESP32-C3-MINI-1\n\u5c4f\u5e55: 1.54\" ST7789 (240\u00d7240)')
run.font.size = Pt(12)

doc.add_page_break()

# TOC
doc.add_heading('\u76ee\u5f55', level=1)
toc = [
    '1. \u7cfb\u7edf\u6982\u8ff0',
    '2. \u786c\u4ef6\u6846\u56fe',
    '3. \u6838\u5fc3\u5668\u4ef6\u9009\u578b',
    '4. IO\u5f15\u811a\u5206\u914d\u8868',
    '5. \u7535\u6e90\u8bbe\u8ba1',
    '6. LCD\u63a5\u53e3\u7535\u8def',
    '7. DHT11\u6e29\u6e7f\u5ea6\u4f20\u611f\u5668',
    '8. USB\u63a5\u53e3\u7535\u8def',
    '9. Strapping\u5f15\u811a\u914d\u7f6e\uff08\u91cd\u8981\uff01\uff09',
    '10. PCB Layout\u6ce8\u610f\u4e8b\u9879',
    '11. \u710a\u63a5\u9a8c\u8bc1\u6e05\u5355',
    '12. BOM\u7269\u6599\u6e05\u5355',
]
for item in toc:
    doc.add_paragraph(item)

doc.add_page_break()

# 1
doc.add_heading('1. \u7cfb\u7edf\u6982\u8ff0', level=1)
doc.add_paragraph(
    'PC Monitor\u662f\u4e00\u6b3ePC\u526f\u5c4f\u76d1\u63a7\u8bbe\u5907\uff0c\u901a\u8fc7ESP32-C3\u8fde\u63a5WiFi\u540e\u63a5\u6536PC\u7aef\u53d1\u9001\u7684\u7cfb\u7edf\u6570\u636e\uff0c'
    '\u5b9e\u65f6\u663e\u793a\u57281.54\u5bf8TFT LCD\u5c4f\u5e55\u4e0a\u3002\u8bbe\u5907\u540c\u65f6\u96c6\u6210DHT11\u4f20\u611f\u5668\uff0c\u53ef\u76d1\u6d4b\u5ba4\u5185\u6e29\u6e7f\u5ea6\u3002'
)
doc.add_paragraph('\u4e3b\u8981\u529f\u80fd\uff1a')
for item in [
    '\u5b9e\u65f6\u663e\u793aCPU\u5360\u7528\u7387\u3001\u6e29\u5ea6',
    '\u5b9e\u65f6\u663e\u793a\u5185\u5b58\u4f7f\u7528\u7387',
    '\u5b9e\u65f6\u663e\u793a\u7f51\u7edc\u4e0a\u884c/\u4e0b\u884c\u901f\u5ea6\uff08\u9700PC\u8fde\u63a5\uff09',
    '\u5ba4\u5185\u6e29\u6e7f\u5ea6\u76d1\u6d4b\uff08DHT11\uff09',
    'WiFi\u8054\u7f51\u81ea\u52a8\u8fde\u63a5PC\u670d\u52a1\u7aef',
    '\u901a\u8fc7USB-C\u4f9b\u7535\uff0c\u540c\u65f6\u652f\u6301\u56fa\u4ef6\u70e7\u5f55',
]:
    doc.add_paragraph(item, style='List Bullet')

# 2
doc.add_heading('2. \u786c\u4ef6\u6846\u56fe', level=1)
for item in [
    '\u4e3b\u63a7\uff1aESP32-C3-MINI-1 \u6a21\u5757',
    '\u663e\u793a\uff1a1.54\u5bf8 ST7789 SPI TFT LCD (240\u00d7240)',
    '\u4f20\u611f\u5668\uff1aDHT11 \u6570\u5b57\u6e29\u6e7f\u5ea6\u4f20\u611f\u5668',
    '\u7535\u6e90\uff1aUSB-C \u2192 SS34\u9632\u53cd\u63a5 \u2192 ME6211C33M5G LDO \u2192 3.3V',
    '\u70e7\u5f55\uff1aUSB-C\u76f4\u8fde\u6a21\u5757USB D-/D+',
]:
    doc.add_paragraph(item, style='List Bullet')

p = doc.add_paragraph()
run = p.add_run(
    'USB-C(5V) \u2192 SS34 \u2192 ME6211(3.3V) \u2192 ESP32-C3-MINI-1\n'
    '                                    \u21d5 SPI\n'
    '                              1.54\" ST7789 LCD\n'
    '                                    \u21d5 Single Bus\n'
    '                               DHT11 \u4f20\u611f\u5668'
)
run.font.name = 'Consolas'
run.font.size = Pt(9)

# 3
doc.add_heading('3. \u6838\u5fc3\u5668\u4ef6\u9009\u578b', level=1)

doc.add_heading('3.1 \u4e3b\u63a7\u6a21\u5757', level=2)
doc.add_paragraph('\u578b\u53f7\uff1aESP32-C3-MINI-1')
for item in [
    '\u5185\u6838\uff1aRISC-V 32\u4f4d\u5355\u6838\uff0c\u6700\u9ad8160MHz',
    'Flash\uff1a4MB\uff08\u5185\u7f6e\uff09',
    'RAM\uff1a400KB SRAM',
    'WiFi\uff1a802.11 b/g/n',
    '\u5c01\u88c5\uff1acastellated holes, 40pin',
    '\u5de5\u4f5c\u7535\u538b\uff1a3.0V~3.6V\uff08\u5178\u578b3.3V\uff09',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('3.2 \u663e\u793a\u5c4f', level=2)
doc.add_paragraph('\u578b\u53f7\uff1a1.54\u5bf8 TFT LCD (ST7789\u9a71\u52a8)')
for item in [
    '\u5206\u8fa8\u7387\uff1a240\u00d7240 RGB',
    '\u63a5\u53e3\uff1a4\u7ebfSPI',
    '\u9a71\u52a8IC\uff1aST7789V',
    '\u5de5\u4f5c\u7535\u538b\uff1a2.8V~3.3V',
    '\u7279\u70b9\uff1a\u5168\u89c6\u89d2\u3001\u9ad8\u4eae\u5ea6',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('3.3 \u7535\u6e90\u82af\u7247', level=2)
doc.add_paragraph('\u578b\u53f7\uff1aME6211C33M5G')
for item in [
    '\u7c7b\u578b\uff1a\u4f4e\u538b\u5dee\u7ebf\u6027\u7a33\u538b\u5668 (LDO)',
    '\u8f93\u51fa\u7535\u538b\uff1a3.3V',
    '\u8f93\u51fa\u7535\u6d41\uff1a500mA\uff08\u6700\u5927\uff09',
    '\u538b\u5dee\uff1a200mV @ 100mA',
    '\u5c01\u88c5\uff1aSOT-23-5',
]:
    doc.add_paragraph(item, style='List Bullet')

# 4 - IO Table
doc.add_heading('4. IO\u5f15\u811a\u5206\u914d\u8868', level=1)

table = doc.add_table(rows=1, cols=5)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = '\u6a21\u5757\u5f15\u811a#'
hdr[1].text = '\u4fe1\u53f7'
hdr[2].text = 'GPIO'
hdr[3].text = '\u65b9\u5411'
hdr[4].text = '\u8fde\u63a5\u8bbe\u5907'

io = [
    ('2', 'GPIO2', 'GPIO2', 'OUT', 'LCD_BL \u80cc\u5149\u63a7\u5236'),
    ('3', 'GPIO3', 'GPIO3', 'IN/OUT', 'DHT11 DATA'),
    ('4', 'GPIO4', 'GPIO4', 'OUT', 'LCD_DC \u6570\u636e/\u547d\u4ee4'),
    ('5', 'GPIO5', 'GPIO5', 'OUT', 'LCD_CS SPI\u7247\u9009'),
    ('6', 'GPIO6', 'GPIO6', 'OUT', 'LCD_MOSI SPI\u6570\u636e\u7ebf'),
    ('7', 'GPIO7', 'GPIO7', 'OUT', 'LCD_SCK SPI\u65f6\u949f\u7ebf'),
    ('8', 'GPIO8', 'GPIO8', '\u56fa\u5b9a', '4.7K\u03a9\u4e0a\u62c9\u21923.3V\uff08strap\u5fc5\u987b\uff01\uff09'),
    ('9', 'GPIO9', 'GPIO9', '\u56fa\u5b9a', '4.7K\u03a9\u4e0b\u62c9\u2192GND\uff08strap\u5fc5\u987b\uff01\uff09'),
    ('14', 'GPIO21', 'GPIO21', 'OUT', 'LCD_RST \u663e\u793a\u5c4f\u590d\u4f4d'),
    ('19', '3.3V', '-', 'PWR', 'LDO\u8f93\u51fa3.3V'),
    ('20', 'EN', '-', 'IN', '\u590d\u4f4d\u811a: 10K\u03a9\u4e0a\u62c9+0.1\u03bcF'),
    ('21', '3.3V', '-', 'PWR', 'LDO\u8f93\u51fa3.3V'),
    ('22', 'GND', '-', 'PWR', '\u7535\u6e90\u5730'),
    ('23', 'USB_D-', '-', 'IO', 'USB-C D-\uff08\u76f4\u8fde\uff09'),
    ('24', 'USB_D+', '-', 'IO', 'USB-C D+\uff08\u76f4\u8fde\uff09'),
    ('25', 'GND', '-', 'PWR', '\u7535\u6e90\u5730'),
    ('28', 'GPIO10', 'GPIO10', 'NC', '\u60ac\u7a7a\u4fdd\u7559'),
    ('29', 'GPIO11', 'GPIO11', 'NC', '\u60ac\u7a7a\u4fdd\u7559\uff08\u672a\u6765SDIO\uff09'),
    ('30', 'GPIO12', 'GPIO12', 'NC', '\u60ac\u7a7a\u4fdd\u7559\uff08\u672a\u6765I2C SDA\uff09'),
    ('31', 'GPIO13', 'GPIO13', 'NC', '\u60ac\u7a7a\u4fdd\u7559\uff08\u672a\u6765I2C SCL\uff09'),
]
for r in io:
    row = table.add_row().cells
    for i, v in enumerate(r):
        row[i].text = v

# 5 - Power
doc.add_heading('5. \u7535\u6e90\u8bbe\u8ba1', level=1)
doc.add_heading('5.1 \u7535\u6e90\u67b6\u6784', level=2)
doc.add_paragraph(
    '\u7cfb\u7edf\u91c7\u7528USB-C 5V\u4f9b\u7535\uff0c\u7ecf\u8fc7SS34\u82cd\u7279\u57fa\u4e8c\u6781\u7ba1\u9632\u53cd\u63a5\u4fdd\u62a4\uff0c'
    '\u518d\u7531ME6211C33M5G LDO\u7a33\u538b\u81f33.3V\u4f9b\u5404\u6a21\u5757\u4f7f\u7528\u3002'
)

doc.add_heading('5.2 \u7535\u6e90\u7535\u8def', level=2)
p = doc.add_paragraph()
run = p.add_run(
    'USB-C VBUS(5V) \u2192 SS34\u9633\u6781\n'
    '                  \u4e18\u6781\u2192 ME6211 IN\n'
    '                         \u2228\n'
    '                    C1 10\u03bcF \u2192 GND\n'
    '                    C2 0.1\u03bcF \u2192 GND\n'
    '                         \u2228\n'
    '                  ME6211 OUT \u2192 3.3V\n'
    '                         \u2228\n'
    '                    C3 10\u03bcF \u2192 GND\n'
    '                    C4 0.1\u03bcF \u2192 GND\n'
    '                         \u2228\n'
    '       ESP32-C3 + LCD + DHT11'
)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_heading('5.3 \u7535\u6e90\u8981\u6c42', level=2)
for item in [
    '\u8f93\u5165\u7535\u538b\uff1a5V \u00b110%',
    '\u603b\u7535\u6d41\uff1a\u7ea6200mA\uff08\u5178\u578b\uff09',
    'LDO\u8f93\u51fa\u7535\u6d41\uff1a500mA\u8db3\u591f',
    '3.3V\u8d70\u7ebf\u5bbd\u5ea6\uff1a\u22650.5mm',
    '\u7535\u5bb9\u5c3d\u91cf\u9760\u8fd1LDO\u5f15\u811a\u653e\u7f6e',
]:
    doc.add_paragraph(item, style='List Bullet')

# 6 - LCD
doc.add_heading('6. LCD\u63a5\u53e3\u7535\u8def', level=1)
table = doc.add_table(rows=1, cols=4)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'LCD\u5f15\u811a'
hdr[1].text = 'ESP32 GPIO'
hdr[2].text = '\u529f\u80fd'
hdr[3].text = '\u5907\u6ce8'
lcd = [
    ('RST', 'GPIO21', '\u590d\u4f4d', '\u4f4e\u7535\u5e73\u590d\u4f4d'),
    ('DC', 'GPIO4', '\u6570\u636e/\u547d\u4ee4', '\u9ad8=\u6570\u636e\uff0c\u4f4e=\u547d\u4ee4'),
    ('MOSI', 'GPIO6', 'SPI\u6570\u636e\u7ebf', '\u4e3b\u51fa\u4ece\u5165'),
    ('SCK', 'GPIO7', 'SPI\u65f6\u949f\u7ebf', '\u6a21\u5f0f0'),
    ('CS', 'GPIO5', '\u7247\u9009', '\u4f4e\u7535\u5e73\u6709\u6548'),
    ('BL', 'GPIO2', '\u80cc\u5149', '\u9ad8\u7535\u5e73\u5f00\u542f'),
]
for r in lcd:
    row = table.add_row().cells
    for i, v in enumerate(r):
        row[i].text = v

# 7 - DHT11
doc.add_heading('7. DHT11\u6e29\u6e7f\u5ea6\u4f20\u611f\u5668', level=1)
table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'DHT11\u5f15\u811a'
hdr[1].text = '\u8fde\u63a5'
hdr[2].text = '\u5907\u6ce8'
dht = [
    ('VCC', '3.3V', '3.0~5.5V'),
    ('DATA', 'GPIO3 + 4.7K\u03a9\u21913.3V', '\u5f00\u6dec\u8f93\u51fa\uff0c\u5fc5\u987b\u4e0a\u62c9\uff01'),
    ('GND', 'GND', ''),
]
for r in dht:
    row = table.add_row().cells
    for i, v in enumerate(r):
        row[i].text = v

# 8 - USB
doc.add_heading('8. USB\u63a5\u53e3\u7535\u8def', level=1)
table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'USB-C\u5f15\u811a'
hdr[1].text = '\u8fde\u63a5'
hdr[2].text = '\u5907\u6ce8'
usb = [
    ('VBUS', 'SS34\u9633\u6781', '5V\u8f93\u5165'),
    ('D-', 'ESP32 USB_D-', '\u76f4\u8fde\uff0c\u4e0d\u4ea4\u5dee'),
    ('D+', 'ESP32 USB_D+', '\u76f4\u8fde\uff0c\u4e0d\u4ea4\u5dee'),
    ('GND', 'GND', ''),
    ('CC1/CC2', '5.1K\u03a9\u2192GND', 'USB-C\u8bc6\u522b\u7535\u963b'),
]
for r in usb:
    row = table.add_row().cells
    for i, v in enumerate(r):
        row[i].text = v

# 9 - Strapping
doc.add_heading('9. Strapping\u5f15\u811a\u914d\u7f6e\uff08\u91cd\u8981\uff01\uff09', level=1)
doc.add_paragraph(
    'ESP32-C3\u6709strapping\u5f15\u811a\uff0c\u4e0a\u7535\u65f6\u5fc5\u987b\u4e3a\u7279\u5b9a\u7535\u5e73\uff0c'
    '\u5426\u5219\u6a21\u5757\u65e0\u6cd5\u6b63\u5e38\u542f\u52a8\uff01'
)
table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = '\u5f15\u811a'
hdr[1].text = '\u914d\u7f6e'
hdr[2].text = '\u8bf4\u660e'
strap = [
    ('GPIO8', '4.7K\u03a9\u4e0a\u62c9\u21923.3V', '\u51b3\u5b9a\u4eceFlash\u542f\u52a8\uff0c\u5fc5\u987b\uff01'),
    ('GPIO9', '4.7K\u03a9\u4e0b\u62c9\u2192GND', '\u9009\u62e9\u666e\u901a\u542f\u52a8\uff0c\u5fc5\u987b\uff01'),
]
for r in strap:
    row = table.add_row().cells
    for i, v in enumerate(r):
        row[i].text = v

# 10 - PCB Layout
doc.add_heading('10. PCB Layout\u6ce8\u610f\u4e8b\u9879', level=1)
items = [
    ('SPI\u8d70\u7ebf\uff08MOSI/SCK\uff09', '\u5c3d\u91cf\u77ed\uff08<5cm\uff09\uff0c\u8fdc\u79bbDHT11 DATA\u7ebf\uff0c\u4e0d\u5e76\u884c'),
    ('USB\u5dee\u5206\u5bf9\uff08D-/D+\uff09', '\u7b49\u957f\u8d70\u7ebf\uff0c\u5dee\u5206\u963b\u629790\u03a9\uff0c\u4e0d\u4ea4\u5dee'),
    ('\u7535\u6e90', 'LDO\u9760\u8fd1USB\u5ea7\uff0c\u7535\u5bb9\u9760\u8fd1LDO\u5f15\u811a\uff0c3.3V\u8d70\u7ebf\u22650.5mm'),
    ('GPIO2\uff08BL\u80cc\u5149\uff09', '\u4e0a\u7535\u9ed8\u8ba4\u9ad8\u7535\u5e73\uff0c\u53ef\u52a0NMOS\u6216\u4e32\u8054100\u03a9\u963b\u63a7\u5236'),
    ('\u5929\u7ebf', '\u6a21\u5757\u5e95\u90e8\u5929\u7ebf\u533a\u57df\u4e0d\u94fa\u94dc\uff0c\u4e0d\u8986\u5730\uff0c\u671d\u5411PCB\u8fb9\u7f18'),
    ('\u6a21\u5757\u5b89\u88c5', '\u6ce8\u610fPin1\u65b9\u5411\uff0c\u5efa\u8bae\u70ed\u98ce\u67aa\u6216\u56de\u6d41\u710a'),
]
for title, desc in items:
    p = doc.add_paragraph()
    run = p.add_run(f'\u25a0 {title}\uff1a')
    run.bold = True
    p.add_run(desc)

# 11 - Check list
doc.add_heading('11. \u710a\u63a5\u9a8c\u8bc1\u6e05\u5355', level=1)
checks = [
    '3.3V\u548cGND\u4e0d\u77ed\u8def',
    'GPIO8 \u2192 4.7K\u03a9 \u2192 3.3V\uff08\u5426\u5219\u4e0d\u542f\u52a8\uff01\uff09',
    'GPIO9 \u2192 4.7K\u03a9 \u2192 GND\uff08\u5426\u5219\u4e0d\u542f\u52a8\uff01\uff09',
    'EN: 10K\u03a9\u4e0a\u62c9\u21923.3V + 0.1\u03bcF\u2192GND',
    'DHT11 DATA: 4.7K\u03a9\u4e0a\u62c9\u21923.3V',
    'LCD Pin1\u5bf9\u9f50\uff08\u767d\u8272\u6807\u8bb0\u6216\u65b9\u5f62\u710a\u76d8\uff09',
    'USB D-/D+ \u76f4\u8fde\uff0c\u65e0\u4ea4\u5dee',
    'SS34\u65b9\u5411\u6b63\u786e\uff08\u4e1d\u5370\u7bad\u5934\u6307\u5411LDO\uff09',
    'LDO\u7535\u5bb9\u5df2\u710a\u63a5\uff0810\u03bcF+0.1\u03bcF\u5404\u4e24\u7ec4\uff09',
    '\u6a21\u5757\u5929\u7ebf\u533a\u57df\u4e0d\u94fa\u94dc',
    'USB-C CC1/CC2: 5.1K\u03a9\u2192GND',
]
for item in checks:
    doc.add_paragraph(f'\u2610  {item}')

# 12 - BOM
doc.add_heading('12. BOM\u7269\u6599\u6e05\u5355', level=1)
table = doc.add_table(rows=1, cols=4)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = '\u5668\u4ef6'
hdr[1].text = '\u6570\u91cf'
hdr[2].text = '\u5355\u4ef7(\u00a5)'
hdr[3].text = '\u5907\u6ce8'
bom = [
    ('ESP32-C3-MINI-1', '1', '8.0', '\u6dd8\u5b9d/\u7acb\u521b'),
    ('1.54\" ST7789 LCD', '1', '12.0', '\u542bFPC+\u8f6c\u63a5\u677f'),
    ('ME6211C33M5G', '1', '0.5', 'SOT-23-5'),
    ('SS34', '1', '0.2', 'SMA\u5c01\u88c5'),
    ('USB-C\u6bcd\u5ea7 16P', '1', '1.0', '\u7acb\u521b\u5546\u57ce'),
    ('4.7K\u03a9 0805', '3', '0.03', 'GPIO8/9/DHT11'),
    ('10K\u03a9 0805', '1', '0.03', 'EN\u4e0a\u62c9'),
    ('5.1K\u03a9 0805', '2', '0.03', 'USB-C CC'),
    ('10\u03bcF 0805', '2', '0.1', 'LDO\u8f93\u5165\u8f93\u51fa'),
    ('0.1\u03bcF 0805', '2', '0.05', 'EN+LDO\u6ee4\u6ce2'),
    ('PCB\u6253\u6837 5\u00d78cm', '1', '2.0', '\u7acb\u521b5\u7247/\u00a510'),
    ('DHT11\u6a21\u5757', '1', '2.0', '\u6216DHT22=\u00a58'),
]
for r in bom:
    row = table.add_row().cells
    for i, v in enumerate(r):
        row[i].text = v

p = doc.add_paragraph()
run = p.add_run('\u603b\u8ba1(1.54\u5bf8\u65b9\u6848)\uff1a\u00a526~30/\u5957\uff08\u4e0d\u542bSMT\u710a\u63a5\u548c\u5305\u88c5\uff09')
run.bold = True

doc.save('E:/CODEX-Pj/ESP32/PCB\u8bbe\u8ba1/PC_Monitor_\u786c\u4ef6\u539f\u7406\u56fe\u8bbe\u8ba1\u53c2\u8003\u6307\u5357.docx')
print('OK')
